# CREATE WEEKLY REPEAT VIA CRON OR WINDOWS SCHEDULER

import json
import requests
import sqlite3 as sql
import pandas as pd
import docx # pip install python-docx
import asyncio
import nest_asyncio
import json
import aiohttp
from understat import Understat
import difflib
from fuzzywuzzy import fuzz

pd.set_option('precision', 3)
pd.options.display.precision = 1

# Checks name length and adjusts player_name if it's more than 2 words or has other errors
def checkNameLength(df):
    string = df['player_name']
    
    names = string.split(' ')
    web_names = df['web_name'].split(' ')
    count = len(names)
    if count > 2:
        # If first name is web name, don't repeat
        if names[0] == web_names[-1]:
            string = names[0]
        else:
            string = names[0] + ' ' + web_names[-1]
    if count == 2:
        # Ensure player name not repeated
        if names[0] == names[-1]:
            string = web_names[0] + ' ' + web_names[-1]
    return string

### Get FPL Data ###
# Connects to the API and converts to json
url = requests.get('https://fantasy.premierleague.com/api/bootstrap-static/')
data = url.json()

# Clean version of the data that is readable if needed
data_clean = json.dumps(data, indent = 2)
parsed_data = json.loads(data_clean)

# Checks which game week the API is updated to
gameweeks = data['events']
for gameweek in gameweeks:
    if gameweek['is_current'] == True:
        week = gameweek['name']
        break
    else:
        continue 

# Creates the database name to connect to depending on what Gameweek it is
week = week.replace(' ', '')
season = '2020'

# Saves json clean format to word doc for reading
jsonclean = 'FPL-{}-{}.docx'.format(season, week)
doc = docx.Document()
doc.add_paragraph(data_clean)
doc.save(jsonclean)

### Get Understat Data ###

nest_asyncio.apply()

ustat = []

# Connects to understat and retrieves all data from get_league_players
async def main():
    async with aiohttp.ClientSession() as session:
        understat = Understat(session)
        player = await understat.get_league_players(
            "epl", 2020,
        )
        #print(json.dumps(players))
        ustat.append(json.dumps(player))

loop = asyncio.get_event_loop()
loop.run_until_complete(main())

# Save understat json to doc
ustat_jsonclean = 'Understat-{}-{}.docx'.format(season, week)
doc = docx.Document()
doc.add_paragraph(ustat)
doc.save(ustat_jsonclean)

### Convert FPL Data to pandas ### 

players = data['elements']

df = pd.DataFrame(players)

cols = ['first_name', 'second_name', 'web_name', 'id', 'element_type', 'team', 'team_code', 'now_cost', 'total_points', 'goals_scored', 'assists', 'clean_sheets', 'bonus', 'saves', 'yellow_cards', 'red_cards', 'form', 
        'points_per_game', 'penalties_saved', 'penalties_missed', 'influence', 'creativity', 'threat', 'ict_index', 'transfers_in_event', 'transfers_out_event' ]
df = df[cols]

teams = {1: 'Arsenal', 2: 'Aston Villa', 3: 'Brighton', 4: 'Burnley', 5: 'Chelsea', 6: 'Crystal Palace',
7: 'Everton', 8: 'Fulham', 9: 'Leicester', 10: 'Leeds', 11: 'Liverpool', 12: 'Manchester City', 13: 'Manchester United',
14: 'Newcastle United', 15: 'Sheffield United', 16: 'Southampton', 17: 'Tottenham Hotspur', 
18: 'West Brom', 19: 'West Ham', 20: 'Wolves'}

df.loc[:, 'team'] = df['team'].map(teams)

### Convert Understat Data to pandas ### 

ustat = json.loads(ustat[0])
dfu = pd.DataFrame(ustat)

colsu = ['key_passes', 'npg', 'npxG', 'player_name', 'shots', 'xA', 'xG', 'xGBuildup', 'xGChain']
dfu = dfu[colsu]

dfu.loc[:, 'first_name'] = dfu['player_name'].str.split(' ').str[0]
dfu.loc[:, 'second_name'] = dfu['player_name'].str.split(' ').str[-1]

### Merge both dataframes on player_name ### 

df['player_name'] = df['first_name'] + ' ' + df['second_name']

df.loc[:, 'player_name'] = df.apply(lambda x: checkNameLength(x), axis=1)

# Finds missing names in both datasets
names_u = list(set(list(dfu.player_name)) - set(list(df.player_name)))
names =  list(set(list(df.player_name)) - set(list(dfu.player_name)))

# Splits name and finds surname in names_u
matches = []
for name in names:
    sub = name.split(' ')[-1]
    match = [(name, s) for s in names_u if sub in s]
    matches.append(match)
    
# Finds name matches in both datasets by fuzzy string matching
name_change = {}
for match in matches: 
    # If only one match, save as key-value pair
    if len(match) == 1:
        name_change[match[0][0]] = match[0][1]
        matches.remove(match)
    # If multiple matches, take the key-value pair with highest fuzzy match ratio
    elif len(match) > 1:
        prefu = 0
        for m in match:
            fu = fuzz.ratio(m[0], m[1])
            if fu > prefu:
                save = m
                prefu = fu
        name_change[save[1]] = save[0]
        matches.remove(match)

# Manually fix errors in name_change dictionary
name_change['Felipe Anderson Pereira Gomes'] = 'Felipe Anderson'
name_change['Thiago Silva'] = 'Thiago Silva'
name_change['Dele Alli'] = 'Dele Alli'
name_change['Bernard'] = 'Bernard'
name_change['David Martin'] = 'David Martin'

# Change instances of player name to match
dfu.loc[:, 'player_name'].replace(name_change, inplace=True)
df.loc[:, 'player_name'].replace(name_change, inplace=True)

# Check missing names
print(dfu[~dfu['player_name'].isin(df['player_name'].tolist())]['player_name'])

# Manually build dictionary to fill in names that don't match
manual_name_change = {'Jorge Luiz Frello Filho': 'Jorginho',
                      'Aleksandar Mitrović': 'Aleksandar Mitrovic',
                      'Benjamin Chilwell': 'Ben Chilwell',
                      'Bobby Decordova-Reid': 'Bobby Reid',
                      'Rodrigo Moreno': 'Rodrigo',
                      'Romain Saïss': 'Romain Saiss',
                      'Gabriel Magalhães': 'Gabriel',
                      'Nicolas Pépé': 'Nicolas Pepe',
                      'Tanguy NDombele Alvaro': 'Tanguy Ndombele',
                      'Donny Beek': 'Donny van de Beek',
                      'Joelinton Joelinton': 'Joelinton',
                      'Thiago Alcántara': 'Thiago',
                      'Alexis Allister': 'Alexis Mac Allister',
                      'Bamidele Alli': 'Dele Alli',
                      'Ahmed Mohamady': 'Ahmed Elmohamady',
                      'Rodrigo Hernandez': 'Rodri',
                      'Ahmed Hegazy': 'Ahmed Hegazi',
                      'Emiliano Martínez': 'Emiliano Martinez',
                      'Kepa Arrizabalaga': 'Kepa',
                      'Çaglar Söyüncü': 'Caglar Söyüncü',
                      'Franck Zambo': 'André-Frank Anguissa',
                      'Bernard Bernard': 'Bernard',
                      'Jack O&#039;Connell': 'Jack O\'Connell',
                      'Daniel N&#039;Lundulu': 'Daniel N\'Lundulu',
                      'Dara O&#039;Shea': 'Dara O\'Shea',
                      'Vitor Ferreira': 'Vitinha',
                      'N&#039;Golo Kanté': 'N\'Golo Kanté',
                      'Fernando Fernandinho': 'Fernandinho',
                      'Saïd Benrahma': 'Said Benrahma'
                     }

# Replace names via manual_name_change dictionary
dfu.loc[:, 'player_name'].replace(manual_name_change, inplace=True)
df.loc[:, 'player_name'].replace(manual_name_change, inplace=True)

# Check to ensure no names missing
print(dfu[~dfu['player_name'].isin(df['player_name'].tolist())]['player_name'])

# Join DataFrames
df = pd.concat([df.set_index('player_name'), dfu.set_index('player_name')], axis=1, join='outer').reset_index()
# Drop unnecessary columns
df.drop(['first_name', 'team_code', 'second_name', 'first_name', 'second_name'], axis=1, inplace=True)
# Change Element type bto position
position_map = {1: 'GK', 2: 'DEF', 3: 'MID', 4: 'FWD'}
df.loc[:, 'element_type'] = df['element_type'].map(position_map)
df.rename(columns={'element_type':'position', 'index': 'player_name'}, inplace=True)
# Fill na to prevent NULL error in SQL
df.fillna(0, inplace=True)
# Save to csv
df.to_csv('FPL-{}-{}.csv'.format(season, week))

### Convert to SQL
# db name
db = 'EPL-Data-{}.sqlite'.format(season)

conn = sql.connect(db)
cur = conn.cursor()

# Create table that specifies gameweek and season
sql_create_player_table = '''
CREATE TABLE IF NOT EXISTS Player (
id    integer    PRIMARY KEY,
name  text       NOT NULL,
web_name    text    NOT NULL,
player_id    integer    NOT NULL,
position   text   NOT NULL,
team   text   NOT NULL,
value    real    NOT NULL,
total_points    integer    NOT NULL,
goals_scored    integer    NOT NULL,
assists    integer    NOT NULL,
clean_sheets    integer    NOT NULL,
bonus    integer    NOT NULL,
saves    integer     NOT NULL,
yellow_cards    integer    NOT NULL,
red_cards    integer    NOT NULL,
form     real    NOT NULL,
points_per_game    real    NOT NULL,
penalties_saved    integer    NOT NULL,
penalties_missed    integer    NOT NULL,
influence     real    NOT NULL,
creativity    real    NOT NULL,
threat    real    NOT NULL,
ict_index   real    NOT NULL,
transfers_in_event    integer    NOT NULL,
transfers_out_event    integer    NOT NULL,
key_passes    integer    NOT NULL,
npg    real    NOT NULL,
npxG    real    NOT NULL,
shots    integer    NOT NULL,
xA    real    NOT NULL,
xG    real    NOT NULL, 
xGBuildup    real    NOT NULL,
xGChain    real    NOT NULL,
gameweek    integer    NOT NULL
);'''


# Create SQL table in database
cur.execute(sql_create_player_table)
conn.commit()

df.loc[:, 'gameweek'] = week

# Iterates over each row and inserts into db. All in same table.
for i, player in df.iterrows():
    name = player['player_name']
    web_name = player['web_name']
    player_id = player['id']
    position = player['position']
    team = player['team']
    cost = player['now_cost']
    total_points = player['total_points']
    goals_scored = player['goals_scored']
    assists = player['assists']
    clean_sheets = player['clean_sheets']
    bonus = player['bonus']
    saves = player['saves']
    yellow_cards = player['yellow_cards']
    red_cards = player['red_cards']
    form = player['form']
    points_per_game = player['points_per_game']
    penalties_saved = player['penalties_saved']
    penalties_missed = player['penalties_missed']
    influence = player['influence']
    creativity = player['creativity']
    threat = player['threat']
    ict_index = player['ict_index']
    transfers_in_event = player['transfers_in_event']
    transfers_out_event = player['transfers_out_event']
    key_passes = player['key_passes']
    npg = player['npg']
    npxG = player['npxG']
    shots = player['shots']
    xA = player['xA']
    xG = player['xG']
    xGBuildup = player['xGBuildup']
    xGChain = player['xGChain']
    gameweek = week
    
    cur.execute('''INSERT OR REPLACE INTO PLAYER
    (name, web_name, player_id, position, team, value, total_points, goals_scored, assists, clean_sheets,
    bonus, saves, yellow_cards, red_cards, form, points_per_game, penalties_saved, penalties_missed,
    influence, creativity, threat, ict_index, transfers_in_event, transfers_out_event, key_passes,
    npg, npxG, shots, xA, xG, xGBuildup, xGChain, gameweek) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
    ''', (name, web_name, player_id, position, team, cost, total_points, goals_scored, assists, clean_sheets,
    bonus, saves, yellow_cards, red_cards, form, points_per_game, penalties_saved, penalties_missed,
    influence, creativity, threat, ict_index, transfers_in_event, transfers_out_event, key_passes,
    npg, npxG, shots, xA, xG, xGBuildup, xGChain, gameweek) )

conn.commit()
